from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel
from typing import Optional, Literal
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
import pdfplumber
from docx import Document
import pandas as pd
import json
import uuid
import os
import sys
import shutil
import secrets
from dotenv import load_dotenv
from detector import detect_columns_with_llm, build_detection_result
from database import init_db, get_db, save_mapping, get_mapping, save_upload, get_uploads
from cleaner import clean_dataframe

# Load environment variables from .env file
load_dotenv()
# SECRET_KEY = os.getenv("SECRET_KEY")
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND_DIR = os.path.join(BASE_DIR, "backend")
if BACKEND_DIR not in sys.path:
    sys.path.append(BACKEND_DIR)

# Initialize FastAPI app
app = FastAPI(title="AuditAI API Running!", debug=True)
# Add CORS middleware to allow frontend to communicate with backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)
# Initialize database tables on startup. Runs once when the app starts
@app.on_event("startup")
async def startup_event():
    init_db()

# Security config for JWT token generation and password hashing
SECRET_KEY = "auditiq-secret-key-change-in-production"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_HOURS = 8
pwd_context = CryptContext(schemes=["sha256_crypt"], deprecated="auto")

# Create uploads directory if it doesn't exist. Define allowed file extensions
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls", "pdf", "docx"}

# Hash a plain text password before storing it in the database
def hash_password(password: str):
    return pwd_context.hash(password)

# Verify a plain text password against a stored hashed password
def verify_password(plain: str, hashed: str):
    return pwd_context.verify(plain, hashed)

# Create a JWT token with an expiry time for a logged in user
def create_token(data: dict):
    expire = datetime.utcnow() + timedelta(hours=ACCESS_TOKEN_EXPIRE_HOURS)
    data.update({"exp": expire})
    return jwt.encode(data, SECRET_KEY, algorithm=ALGORITHM)

# Get file extension from filename
def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

# Extract tables and text from PDF files using pdfplumber. Falls back to raw text if no tables found
def extract_pdf(file_path: str):
    tables = []
    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    headers = table[0]
                    rows = table[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    tables.append(df)
            full_text += page.extract_text() or ""
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Extract tables and paragraphs from DOCX files using python-docx. Falls back to paragraphs if no tables found
def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows, columns=headers)
        tables.append(df)
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Read any supported uploaded file into a DataFrame based on its extension
def read_file_to_df(save_path: str, ext: str):
    if ext == "csv":
        return pd.read_csv(save_path)
    elif ext in ["xlsx", "xls"]:
        return pd.read_excel(save_path)
    elif ext == "pdf":
        df, _ = extract_pdf(save_path)
        return df
    elif ext == "docx":
        df, _ = extract_docx(save_path)
        return df
    return None

# Calculate fill rate per column. Fill rate is the percentage of rows that have a value (0.0 to 1.0)
def calculate_fill_rates(df: pd.DataFrame) -> dict:
    fill_rates = {}
    total = len(df)
    for col in df.columns:
        filled = df[col].replace("", float("nan")).dropna().count()
        fill_rates[col] = round(filled / total, 2) if total > 0 else 0.0
    return fill_rates

# Models for request validation. Pydantic models define the shape of data coming into each endpoint
class Client(BaseModel):
    company_name: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    industry: Optional[str] = None
    address: Optional[str] = None
    status: Optional[str] = "Active"
    kra_pin: Literal[True, False] = False

# User model for creating a new user with a password
class User(BaseModel):
    full_name: str
    email: str
    password: str
    phone: Optional[str] = None
    role: Literal["Admin", "Senior Auditor", "Auditor", "Accountant"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

# UserUpdate model for updating a user without changing their password
class UserUpdate(BaseModel):
    full_name: str
    email: str
    phone: Optional[str] = None
    role: Literal["Admin", "Senior Auditor", "Auditor", "Accountant"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

# LoginRequest model for email and password login
class LoginRequest(BaseModel):
    email: str
    password: str

# PasswordResetRequest model for requesting a password reset token via email
class PasswordResetRequest(BaseModel):
    email: str

# PasswordResetConfirm model for confirming a password reset with a token and new password
class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

# ColumnMapping model for manually saving a column mapping for a client
class ColumnMapping(BaseModel):
    client_id: str
    file_type: Optional[str] = "general"
    original_column: str
    mapped_to: str
    confirmed_by: Optional[str] = None

# Engagement model for creating and updating audit engagements per client
class Engagement(BaseModel):
    client_id: int
    engagement_name: str
    financial_year: str
    status: Optional[str] = "Planning"
    start_date: Optional[str] = None
    end_date: Optional[str] = None

# EngagementTeam model for adding a user to an engagement team with a role
class EngagementTeam(BaseModel):
    engagement_id: int
    user_id: int
    role: str

# AuditSection model for creating and updating audit sections within an engagement
class AuditSection(BaseModel):
    engagement_id: int
    section_name: str
    status: Optional[str] = "Pending"
    assigned_to: Optional[int] = None

# Submission model for creating a submission for an audit section
class Submission(BaseModel):
    engagement_id: int
    section_id: int
    submitted_by: int
    status: Optional[str] = "Draft"
    notes: Optional[str] = None

# SubmissionStatus model for updating the status of a submission
class SubmissionStatus(BaseModel):
    status: Literal["Draft", "Submitted", "Under Review", "Changes Requested", "Approved"]
    notes: Optional[str] = None

# Notification model for sending a notification to a user
class Notification(BaseModel):
    user_id: int
    message: str
    type: Optional[str] = "engagement_alert"

# LLM DETECTION ROUTES  
# Root endpoint to confirm the API is running
@app.get("/")
def root():
    return {"message": "Audit AI API is running"}

# Upload endpoint for AI pipeline. Accepts a file and client_id, saves the file, reads it into a DataFrame.Calculates fill rates and returns a preview of the data for column detection
@app.post("/upload")
async def upload_file_ai(
    file: UploadFile = File(...),
    client_id: str = Form(...)
):
    # Validate file extension against allowed types
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only.")
    # Check file size does not exceed 50MB limit
    MAX_FILE_SIZE = 50
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File size exceeds the maximum limit of {MAX_FILE_SIZE} MB. Uploaded file size: {file_size_mb:.2f} MB.")
    # Reset file pointer after reading for size check, then save to disk
    file.file.seek(0)
    file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    # Handle PDF files, extract tables or text and return fill rates and preview
    if ext == "pdf":
        df, source = extract_pdf(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"PDF uploaded — extracted via {source}"}
    # Handle DOCX files, extract tables or paragraphs and return fill rates and preview
    if ext == "docx":
        df, source = extract_docx(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"DOCX uploaded — extracted via {source}"}
    # Handle CSV and Excel files, read into DataFrame and return fill rates and preview
    try:
        df = pd.read_csv(save_path) if ext == "csv" else pd.read_excel(save_path)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    save_upload(file_id, client_id, file.filename, ext, len(df))
    fill_rates = calculate_fill_rates(df)
    return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": "table",
            "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
            "preview": df.head(5).fillna("").to_dict(orient="records"), "message": "File uploaded and processed successfully"}

# Detect column meanings using LLM. First checks if client has a saved mapping, if yes, skips LLM entirely.If no saved mapping is found, reads the file, extracts sample values and runs LLM detection
@app.post("/detect-columns")
async def detect_columns_endpoint(
    client_id: str = Form(...),
    file_id: str = Form(...),
    columns: str = Form(...),
    file_type: str = Form("general"),
    fill_rates: str = Form("{}")
):
    # Parse columns list from JSON string sent by frontend
    try:
        columns_list = json.loads(columns)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid columns format.")
    # Parse fill rates from JSON string sent by frontend
    try:
        fill_rates_dict = json.loads(fill_rates)
    except json.JSONDecodeError:
        fill_rates_dict = {}
    # Check if client already has a confirmed mapping saved in the database. If all columns are mapped, skip LLM
    saved_mapping = get_mapping(client_id, file_type)
    if saved_mapping:
        all_mapped = all(col in saved_mapping for col in columns_list)
        if all_mapped:
            filtered_mapping = {col: saved_mapping[col] for col in columns_list}
            result = build_detection_result(columns_list, filtered_mapping)
            result["file_id"] = file_id
            result["source"] = "saved_mapping"
            result["message"] = "Mapping loaded from saved client profile — LLM skipped."
            return result
    # No saved mapping found, locate the uploaded file on disk
    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    # Read file and extract first non-empty sample value per column to give LLM context
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        sample_values = {}
        for col in columns_list:
            if col in df.columns:
                non_empty = df[col].dropna().replace("", float("nan")).dropna()
                sample_values[col] = str(non_empty.iloc[0]) if len(non_empty) > 0 else ""
            else:
                sample_values[col] = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    # Run LLM detection passing column names, sample values and fill rates as context
    try:
        mapping = detect_columns_with_llm(columns_list, sample_values, fill_rates_dict)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")
    result = build_detection_result(columns_list, mapping)
    result["file_id"] = file_id
    result["source"] = "llm_detection"
    return result

# Save a confirmed column mapping for a client so future uploads skip LLM detection
@app.post("/save-mapping")
async def save_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    mapping: str = Form(...),
    confirmed_by: str = Form(None)
):
    # Parse mapping dict from JSON string sent by frontend
    try:
        mapping_dict = json.loads(mapping)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid mapping format.")
    if not mapping_dict:
        raise HTTPException(status_code=400, detail="Mapping cannot be empty. Please provide a valid mapping.")
    save_mapping(client_id, file_type, mapping_dict, confirmed_by)
    return {"client_id": client_id, "file_type": file_type, "columns_saved": len(mapping_dict),
            "message": f"Mapping saved successfully for client {client_id} and file type {file_type}."}

# Retrieve a previously saved column mapping for a client and file type
@app.get("/get-mapping/{client_id}")
async def get_mapping_endpoint(client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        return {"client_id": client_id, "file_type": file_type, "mapping": {},
                "message": "No saved mapping found for this client."}
    return {"client_id": client_id, "file_type": file_type, "mapping": mapping,
            "columns_mapped": len(mapping), "message": "Saved mapping retrieved successfully."}

# Get upload history for a client ordered by most recent first
@app.get("/uploads/{client_id}")
async def get_uploads_endpoint(client_id: str):
    uploads = get_uploads(client_id)
    return {"client_id": client_id, "total_uploads": len(uploads), "uploads": uploads}

# Clean the uploaded file using the confirmed mapping. Requires a saved mapping to exist for the client
@app.post("/clean")
async def clean_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
) -> dict:
    # Check that a confirmed mapping exists before attempting to clean
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client. Please detect the columns and confirm the mapping first.")
    # Locate the uploaded file on disk
    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    # Read file into DataFrame
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    # Run the cleaner using the confirmed mapping and return cleaned data with a validation report
    try:
        cleaned_df, report = clean_dataframe(df, mapping)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleaning failed: {str(e)}")
    return {"file_id": file_id, "client_id": client_id, "file_type": file_type,
            "cleaned_data": cleaned_df.fillna("").to_dict(orient="records"),
            "validation_report": report, "message": "File cleaned successfully."}

# MANAGEMENT ROUTES  
# Create a router for all management routes. This router is registered into the main app at the bottom
audit_router = APIRouter()

# CLIENTS.Get all clients from the database
@audit_router.get("/clients")
def get_clients(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients")
    return cursor.fetchall()

# Get a single client by client_id. Returns 404 if not found
@audit_router.get("/clients/{client_id}")
def get_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients WHERE client_id = %s", (client_id,))
    client = cursor.fetchone()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client

# Create a new client and return the new client_id
@audit_router.post("/clients")
def create_client(c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO clients (company_name, contact_person, email, phone, industry, address, status, kra_pin)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin)
    )
    db.commit()
    return {"client_id": cursor.lastrowid, "message": "Client created"}

# Update an existing client by client_id
@audit_router.put("/clients/{client_id}")
def update_client(client_id: int, c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE clients SET company_name=%s, contact_person=%s, email=%s,
           phone=%s, industry=%s, address=%s, status=%s, kra_pin=%s WHERE client_id=%s""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin, client_id)
    )
    db.commit()
    return {"message": "Client updated"}

# Delete a client by client_id
@audit_router.delete("/clients/{client_id}")
def delete_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM clients WHERE client_id = %s", (client_id,))
    db.commit()
    return {"message": "Client deleted"}

# USERS.Get all users joined with their assigned client company name
@audit_router.get("/users")
def get_users(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT u.user_id, u.full_name, u.email, u.phone, u.role, u.status,
               u.assigned_client_id, u.created_at, c.company_name
        FROM users u
        LEFT JOIN clients c ON u.assigned_client_id = c.client_id
    """)
    return cursor.fetchall()

# Get a single user by user_id. Password hash is removed before returning
@audit_router.get("/users/{user_id}")
def get_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE user_id = %s", (user_id,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.pop("password_hash", None)
    return user

# Create a new user. Password is hashed before storing. Returns 400 if email already exists
@audit_router.post("/users")
def create_user(u: User, db=Depends(get_db)):
    hashed = hash_password(u.password)
    cursor = db.cursor()
    try:
        cursor.execute(
            """INSERT INTO users (full_name, email, password_hash, phone, role, assigned_client_id, status)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (u.full_name, u.email, hashed, u.phone, u.role, u.assigned_client_id, u.status)
        )
        db.commit()
        return {"user_id": cursor.lastrowid, "message": "User created"}
    except Exception:
        raise HTTPException(status_code=400, detail="Email already exists")

# Update an existing user by user_id. Does not change password
@audit_router.put("/users/{user_id}")
def update_user(user_id: int, u: UserUpdate, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE users SET full_name=%s, email=%s, phone=%s,
           role=%s, assigned_client_id=%s, status=%s WHERE user_id=%s""",
        (u.full_name, u.email, u.phone, u.role, u.assigned_client_id, u.status, user_id)
    )
    db.commit()
    return {"message": "User updated"}

# Delete a user by user_id
@audit_router.delete("/users/{user_id}")
def delete_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "User deleted"}

# Assign a user to a specific client by updating their assigned_client_id
@audit_router.put("/users/{user_id}/assign/{client_id}")
def assign_user_to_client(user_id: int, client_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE users SET assigned_client_id=%s WHERE user_id=%s", (client_id, user_id))
    db.commit()
    return {"message": "User assigned to client"}

# AUTH
# Login endpoint. Verifies email and password then returns a JWT token and user info
@audit_router.post("/auth/login")
def login(req: LoginRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    # Return 401 if user not found or password is wrong
    if not user or not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    # Return 403 if user account is not active
    if user["status"] != "Active":
        raise HTTPException(status_code=403, detail="Account is inactive")
    token = create_token({"user_id": user["user_id"], "email": user["email"], "role": user["role"]})
    return {"access_token": token, "token_type": "bearer",
            "user": {"user_id": user["user_id"], "full_name": user["full_name"],
                     "email": user["email"], "role": user["role"]}}

# Request a password reset token. Generates a secure token and stores it in the database
@audit_router.post("/auth/password-reset-request")
def password_reset_request(req: PasswordResetRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    token = secrets.token_urlsafe(32)
    expires_at = datetime.utcnow() + timedelta(hours=1)
    cursor2 = db.cursor()
    cursor2.execute("INSERT INTO password_resets (user_id, token, expires_at) VALUES (%s, %s, %s)",
                    (user["user_id"], token, expires_at))
    db.commit()
    return {"message": "Password reset token generated", "token": token}

# Confirm a password reset. Validates the token, updates the password and deletes the used token
@audit_router.post("/auth/password-reset-confirm")
def password_reset_confirm(req: PasswordResetConfirm, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM password_resets WHERE token = %s AND expires_at > NOW()", (req.token,))
    reset = cursor.fetchone()
    # Return 400 if token is invalid or has expired
    if not reset:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    hashed = hash_password(req.new_password)
    cursor2 = db.cursor()
    cursor2.execute("UPDATE users SET password_hash = %s WHERE user_id = %s", (hashed, reset["user_id"]))
    cursor2.execute("DELETE FROM password_resets WHERE token = %s", (req.token,))
    db.commit()
    return {"message": "Password reset successful"}

# COLUMN MAPPINGS 

# Get all column mappings across all clients
@audit_router.get("/column-mappings")
def get_all_mappings(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings")
    return cursor.fetchall()

# Get all column mappings for a specific client
@audit_router.get("/column-mappings/{client_id}")
def get_client_mappings(client_id: str, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

# Manually create a new column mapping for a client
@audit_router.post("/column-mappings")
def create_mapping(m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO column_mappings (client_id, file_type, original_column, mapped_to, confirmed_by)
           VALUES (%s, %s, %s, %s, %s)""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by)
    )
    db.commit()
    return {"id": cursor.lastrowid, "message": "Column mapping created"}

# Update an existing column mapping by mapping_id
@audit_router.put("/column-mappings/{mapping_id}")
def update_mapping(mapping_id: int, m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE column_mappings SET client_id=%s, file_type=%s, original_column=%s,
           mapped_to=%s, confirmed_by=%s WHERE id=%s""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by, mapping_id)
    )
    db.commit()
    return {"message": "Column mapping updated"}

# Delete a column mapping by mapping_id
@audit_router.delete("/column-mappings/{mapping_id}")
def delete_mapping(mapping_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM column_mappings WHERE id = %s", (mapping_id,))
    db.commit()
    return {"message": "Column mapping deleted"}

# ENGAGEMENTS 
# Get all engagements joined with their client company name, ordered by most recent first
@audit_router.get("/engagements")
def get_engagements(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        ORDER BY e.created_at DESC
    """)
    return cursor.fetchall()

# Get a single engagement by engagement_id. Returns 404 if not found
@audit_router.get("/engagements/{engagement_id}")
def get_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    return engagement

# Create a new engagement and auto-generate default audit sections: Revenue, Expenses, Inventory, Cash & Bank
@audit_router.post("/engagements")
def create_engagement(e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO engagements (client_id, engagement_name, financial_year, status, start_date, end_date)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date)
    )
    engagement_id = cursor.lastrowid
    for section in ["Revenue", "Expenses", "Inventory", "Cash & Bank"]:
        cursor.execute("INSERT INTO audit_sections (engagement_id, section_name) VALUES (%s, %s)", (engagement_id, section))
    db.commit()
    return {"engagement_id": engagement_id, "message": "Engagement created with default audit sections"}

# Update an existing engagement by engagement_id
@audit_router.put("/engagements/{engagement_id}")
def update_engagement(engagement_id: int, e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE engagements SET client_id=%s, engagement_name=%s, financial_year=%s,
           status=%s, start_date=%s, end_date=%s WHERE engagement_id=%s""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date, engagement_id)
    )
    db.commit()
    return {"message": "Engagement updated"}

# Delete an engagement and all its related sections and team members
@audit_router.delete("/engagements/{engagement_id}")
def delete_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagements WHERE engagement_id = %s", (engagement_id,))
    db.commit()
    return {"message": "Engagement deleted"}

# ENGAGEMENT TEAM
# Get all team members for an engagement joined with their user details
@audit_router.get("/engagements/{engagement_id}/team")
def get_engagement_team(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT et.*, u.full_name, u.email, u.role FROM engagement_team et
        LEFT JOIN users u ON et.user_id = u.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

# Add a user to an engagement team with a specific role
@audit_router.post("/engagements/{engagement_id}/team")
def add_team_member(engagement_id: int, t: EngagementTeam, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("INSERT INTO engagement_team (engagement_id, user_id, role) VALUES (%s, %s, %s)",
                   (engagement_id, t.user_id, t.role))
    db.commit()
    return {"team_id": cursor.lastrowid, "message": "Team member added"}

# Remove a user from an engagement team by engagement_id and user_id
@audit_router.delete("/engagements/{engagement_id}/team/{user_id}")
def remove_team_member(engagement_id: int, user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id=%s AND user_id=%s", (engagement_id, user_id))
    db.commit()
    return {"message": "Team member removed"}

# AUDIT SECTIONS ROUTES
# Get all audit sections for an engagement joined with the assigned user's name
@audit_router.get("/engagements/{engagement_id}/sections")
def get_audit_sections(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as assigned_to_name FROM audit_sections s
        LEFT JOIN users u ON s.assigned_to = u.user_id
        WHERE s.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

# Add a new audit section to an engagement
@audit_router.post("/engagements/{engagement_id}/sections")
def add_audit_section(engagement_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("INSERT INTO audit_sections (engagement_id, section_name, status, assigned_to) VALUES (%s, %s, %s, %s)",
                   (engagement_id, s.section_name, s.status, s.assigned_to))
    db.commit()
    return {"section_id": cursor.lastrowid, "message": "Audit section added"}

# Update an existing audit section by section_id
@audit_router.put("/audit-sections/{section_id}")
def update_audit_section(section_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE audit_sections SET section_name=%s, status=%s, assigned_to=%s WHERE section_id=%s",
                   (s.section_name, s.status, s.assigned_to, section_id))
    db.commit()
    return {"message": "Audit section updated"}

# Delete an audit section by section_id
@audit_router.delete("/audit-sections/{section_id}")
def delete_audit_section(section_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE section_id = %s", (section_id,))
    db.commit()
    return {"message": "Audit section deleted"}

# SUBMISSIONS 
# Get all submissions joined with user, engagement and section details, ordered by most recent first
@audit_router.get("/submissions")
def get_all_submissions(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        ORDER BY s.created_at DESC
    """)
    return cursor.fetchall()

# Get a single submission by submission_id. Returns 404 if not found
@audit_router.get("/submissions/{submission_id}")
def get_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.submission_id = %s
    """, (submission_id,))
    submission = cursor.fetchone()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    return submission

# Create a new submission. If status is Submitted, auto-notify Senior Auditors and Auditors on the engagement team
@audit_router.post("/submissions")
def create_submission(s: Submission, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO submissions (engagement_id, section_id, submitted_by, status, notes)
           VALUES (%s, %s, %s, %s, %s)""",
        (s.engagement_id, s.section_id, s.submitted_by, s.status, s.notes)
    )
    submission_id = cursor.lastrowid
    # If submission is marked as Submitted, send notifications to auditors on the team
    if s.status == "Submitted":
        cursor.execute("""
            SELECT e.engagement_name, sec.section_name FROM engagements e
            LEFT JOIN audit_sections sec ON sec.engagement_id = e.engagement_id
            WHERE e.engagement_id = %s AND sec.section_id = %s
        """, (s.engagement_id, s.section_id))
        info = cursor.fetchone()
        if info:
            message = f"{info['section_name']} for {info['engagement_name']} is ready for review"
            cursor.execute("""
                SELECT u.user_id FROM users u
                INNER JOIN engagement_team et ON u.user_id = et.user_id
                WHERE et.engagement_id = %s AND u.role IN ('Senior Auditor', 'Auditor')
            """, (s.engagement_id,))
            for auditor in cursor.fetchall():
                cursor.execute("INSERT INTO notifications (user_id, message, type) VALUES (%s, %s, %s)",
                               (auditor['user_id'], message, 'engagement_alert'))
    db.commit()
    return {"submission_id": submission_id, "message": "Submission created"}

# Update the status of a submission by submission_id
@audit_router.put("/submissions/{submission_id}/status")
def update_submission_status(submission_id: int, s: SubmissionStatus, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE submissions SET status=%s, notes=%s WHERE submission_id=%s",
                   (s.status, s.notes, submission_id))
    db.commit()
    return {"message": f"Submission status updated to {s.status}"}

# Delete a submission by submission_id
@audit_router.delete("/submissions/{submission_id}")
def delete_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM submissions WHERE submission_id = %s", (submission_id,))
    db.commit()
    return {"message": "Submission deleted"}

# NOTIFICATIONS 
# Get all notifications for a user ordered by most recent first
@audit_router.get("/notifications/{user_id}")
def get_user_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
    return cursor.fetchall()

# Get only unread notifications for a user ordered by most recent first
@audit_router.get("/notifications/{user_id}/unread")
def get_unread_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s AND is_read = FALSE ORDER BY created_at DESC", (user_id,))
    return cursor.fetchall()

# Mark a single notification as read by notification_id
@audit_router.put("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE notification_id = %s", (notification_id,))
    db.commit()
    return {"message": "Notification marked as read"}

# Mark all notifications for a user as read
@audit_router.put("/notifications/{user_id}/read-all")
def mark_all_read(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "All notifications marked as read"}

# CLIENT FILE UPLOAD 
# Upload a file for a specific client. Saves the file to disk and records it in the uploads table for audit trail.
@audit_router.post("/clients/{client_id}/upload")
def upload_client_file(client_id: int, file: UploadFile = File(...), db=Depends(get_db)):
    allowed_types = ["xlsx", "xls", "csv", "pdf", "tiff", "tif", "jpg", "jpeg", "png", "xml", "json", "txt"]
    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail="File format not allowed. Accepted: Excel (.xlsx, .xls), CSV (.csv), PDF (.pdf), Scanned (.jpg, .png, .tiff), ERP (.xml, .json, .txt)")
    file_path = f"{UPLOAD_DIR}/{client_id}_{file.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    cursor = db.cursor()
    cursor.execute("INSERT INTO uploads (client_id, file_name, file_type, file_path) VALUES (%s, %s, %s, %s)",
                   (client_id, file.filename, file_ext.upper(), file_path))
    db.commit()
    return {"file_id": cursor.lastrowid, "filename": file.filename, "type": file_ext.upper(), "message": "File uploaded successfully"}

# Get all files uploaded for a specific client
@audit_router.get("/clients/{client_id}/files")
def get_client_files(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM uploads WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

# Get all uploaded files across all clients joined with company name. For admin use only — should be protected in production
@audit_router.get("/files")
def get_all_files(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT f.*, c.company_name FROM uploads f
        LEFT JOIN clients c ON f.client_id = c.client_id
        ORDER BY f.upload_date DESC
    """)
    return cursor.fetchall()

# REGISTER AUDIT ROUTER INTO MAIN APP.Plug all management routes into the main app. All audit_router routes are available on the same server
app.include_router(audit_router)

# Entry point when running the file directly.
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)

