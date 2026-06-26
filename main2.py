from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel
from typing import Optional, Literal
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
import pdfplumber
from docx import Document
import pandas as pd
import json
import uuid
import os
import sys
import shutil
import secrets
import hashlib
from dotenv import load_dotenv
from detector import detect_columns_with_llm, build_detection_result, suggest_file_type
from database import (
    init_db, get_db, save_mapping, get_mapping, save_upload, get_uploads,
    save_cleaning_acknowledgment, get_acknowledged_issue_ids,
    save_cleaning_corrections, get_cleaning_corrections,
    save_fingerprint, get_fingerprint,
    save_cleaning_snapshot, get_cleaning_snapshot,
)
from cleaner import clean_dataframe
from excel_export import build_cleaning_workbook
from excel_diff import diff_uploaded_against_snapshot

load_dotenv()
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND_DIR = os.path.join(BASE_DIR, "backend")
if BACKEND_DIR not in sys.path:
    sys.path.append(BACKEND_DIR)

app = FastAPI(title="AuditAI API Running!", debug=True)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    init_db()

SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_HOURS = 8
pwd_context = CryptContext(schemes=["sha256_crypt"], deprecated="auto")

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls", "pdf", "docx"}

# Internal columns the system adds for its own bookkeeping (row tracking for the
# corrected-Excel diff flow, duplicate-row marking during cleaning) that should
# never be treated as real data and should never reach column detection or mapping.
# These only end up "visible" as real columns if someone uploads an already-exported/
# cleaned workbook through the normal upload flow instead of "Upload Corrected File" —
# the corrected-Excel flow already knows to strip these by name, but the normal
# upload flow has no equivalent protection without this filter.
RESERVED_INTERNAL_COLUMNS = {"_row_id", "_is_duplicate"}

def hash_password(password: str):
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str):
    return pwd_context.verify(plain, hashed)

def create_token(data: dict):
    expire = datetime.utcnow() + timedelta(hours=ACCESS_TOKEN_EXPIRE_HOURS)
    data.update({"exp": expire})
    return jwt.encode(data, SECRET_KEY, algorithm=ALGORITHM)

def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

def extract_pdf(file_path: str):
    tables = []
    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    headers = table[0]
                    rows = table[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    tables.append(df)
            full_text += page.extract_text() or ""
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows, columns=headers)
        tables.append(df)
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

def read_file_to_df(save_path: str, ext: str):
    if ext == "csv":
        df = pd.read_csv(save_path, dtype=str)
    elif ext in ["xlsx", "xls"]:
        df = pd.read_excel(save_path, dtype=str)
    elif ext == "pdf":
        df, _ = extract_pdf(save_path)
        return df
    elif ext == "docx":
        df, _ = extract_docx(save_path)
        return df
    else:
        return None
    if df is not None:
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
        # Drop any reserved internal column (e.g. _row_id from a previously
        # exported workbook re-uploaded by mistake) before it's ever treated as
        # real data — see RESERVED_INTERNAL_COLUMNS for why this exists.
        df = df.drop(columns=[c for c in df.columns if c in RESERVED_INTERNAL_COLUMNS], errors='ignore')
    return df

def calculate_fill_rates(df: pd.DataFrame) -> dict:
    fill_rates = {}
    total = len(df)
    for col in df.columns:
        filled = df[col].replace("", float("nan")).dropna().count()
        fill_rates[col] = round(filled / total, 2) if total > 0 else 0.0
    return fill_rates

def compute_schema_fingerprint(columns: list) -> str:
    sorted_cols = sorted([col.lower().strip() for col in columns])
    fingerprint = hashlib.md5(json.dumps(sorted_cols).encode()).hexdigest()
    return fingerprint

def locate_uploaded_file(file_id: str):
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            return path, extension
    return None, None

def normalize_header_label(value) -> str:
    label = str(value or "").strip()
    if label.startswith("[UNRESOLVED]"):
        label = label.replace("[UNRESOLVED]", "", 1).strip()
    return label

def issue_fingerprint(file_id: str, client_id: str, file_type: str, issue: dict) -> str:
    raw = "|".join([
        str(file_id),
        str(client_id),
        str(file_type),
        str(issue.get("row_index", "")),
        str(issue.get("column", "")),
        str(issue.get("original_value", "")),
        str(issue.get("issue", "")),
    ])
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()

def enrich_issues_with_ids(report: dict, file_id: str, client_id: str, file_type: str) -> dict:
    for issue in report.get("issues", []):
        issue["issue_id"] = issue_fingerprint(file_id, client_id, file_type, issue)
        issue["decision"] = "pending"
    return report

def rebuild_report_counts(report: dict) -> dict:
    issues = report.get("issues", [])
    total_rows = report.get("total_rows", 0)

    row_issues = [
        i for i in issues
        if i.get("row_index") not in ("N/A", None)
    ]

    column_issues = [
        i for i in issues
        if i.get("row_index") in ("N/A", None)
    ]

    flagged_rows = len(set(
    i.get("row_index") for i in row_issues
    if i.get("row_index") not in (None, "N/A")
    ))

    report["flagged_rows"] = flagged_rows
    report["clean_rows"] = total_rows - flagged_rows

    report["row_issues"] = len(row_issues)
    report["column_issues"] = len(column_issues)

    report["total_issues"] = len(issues)
    report["high_issues"] = len([i for i in issues if i.get("severity") == "high"])
    report["medium_issues"] = len([i for i in issues if i.get("severity") == "medium"])
    report["info_issues"] = len([i for i in issues if i.get("severity") == "info"])

    return report

def filter_acknowledged_issues(report: dict, file_id: str, client_id: str, file_type: str) -> dict:
    acknowledged = get_acknowledged_issue_ids(file_id, client_id, file_type)
    if not acknowledged:
        return report
    report["issues"] = [
        issue for issue in report.get("issues", [])
        if issue.get("issue_id") not in acknowledged
    ]
    return rebuild_report_counts(report)

def apply_saved_corrections(df: pd.DataFrame, mapping: dict, corrections: list) -> pd.DataFrame:
    """
    Apply every saved correction to the source dataframe before cleaning runs.
    Handles cell-value corrections, row deletions (column_name == "_row_deleted"),
    and column deletions (column_name starts with "_column_deleted:"). Deletions are
    applied first so they persist on every future cleaning run, not just one response.
    """
    if not corrections:
        return df

    standard_to_original = {
        info.get("mapped_to"): original_col
        for original_col, info in mapping.items()
        if isinstance(info, dict) and info.get("mapped_to") not in ("", "unknown", None)
    }

    rows_to_drop = []
    columns_to_drop = []
    value_corrections = []

    for correction in corrections:
        col = correction.get("column_name", correction.get("column"))
        if col == "_row_deleted":
            rows_to_drop.append(int(correction["row_index"]))
        elif col and str(col).startswith("_column_deleted:"):
            original_col_name = col.split(":", 1)[1]
            columns_to_drop.append(original_col_name)
        else:
            value_corrections.append(correction)

    if rows_to_drop:
        df = df.drop(index=[r for r in rows_to_drop if r in df.index])
    if columns_to_drop:
        # Translate standard mapped names back to original column names
        # because df still has original names at this point (before rename_columns runs)
        resolved_columns_to_drop = []
        for col_name in columns_to_drop:
            # Check if it's already an original column name
            if col_name in df.columns:
                resolved_columns_to_drop.append(col_name)
            else:
                # Look up the original name from the mapping
                original = standard_to_original.get(col_name)
                if original and original in df.columns:
                    resolved_columns_to_drop.append(original)
        df = df.drop(columns=resolved_columns_to_drop)

    for correction in value_corrections:
        row_index = int(correction["row_index"])
        issue_col = correction.get("column_name", correction.get("column"))
        source_col = standard_to_original.get(issue_col, issue_col)
        if source_col in df.columns and row_index in df.index:
            df.at[row_index, source_col] = correction["corrected_value"]
    return df

def adapt_mapping_to_uploaded_headers(mapping: dict, columns: list) -> dict:
    adapted = {}
    column_set = set(columns)
    for original_col, info in mapping.items():
        if not isinstance(info, dict):
            adapted[original_col] = info
            continue
        mapped_to = info.get("mapped_to")
        if original_col in column_set:
            adapted[original_col] = info
        elif mapped_to and mapped_to != "unknown" and mapped_to in column_set:
            adapted[mapped_to] = {**info, "mapped_to": mapped_to}
        else:
            adapted[original_col] = info
    return adapted

def run_cleaning_cycle(file_id: str, client_id: str, file_type: str, mapping: dict):
    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    corrections = get_cleaning_corrections(file_id, client_id, file_type)
    df = apply_saved_corrections(df, mapping, corrections)
    fill_rates = calculate_fill_rates(df)
    try:
        cleaned_df, report = clean_dataframe(df, mapping, fill_rates)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleaning failed: {str(e)}")
    report = enrich_issues_with_ids(report, file_id, client_id, file_type)

    # IMPORTANT: remove acknowledged FIRST
    report = filter_acknowledged_issues(report, file_id, client_id, file_type)

    # recompute counts AFTER filtering
    report = rebuild_report_counts(report)

    # SINGLE SOURCE OF TRUTH for can_proceed: every severity (high, medium, AND
    # info) must be resolved before proceeding. info-severity issues (currently
    # just ambiguous-date warnings) still require an explicit action — either
    # "Correct as-is" (acknowledge-issue, filtered out above) or an actual cell
    # edit (which changes original_value, so the issue won't reappear on the
    # next clean cycle) — they are NOT silently exempt. This prevents an
    # auditor from clearing only the high/medium issues and proceeding without
    # ever having looked at an ambiguous date. Every endpoint below reads this
    # single value instead of recomputing it differently.
    report["can_proceed"] = report["total_issues"] == 0

    return cleaned_df, report

# ── MODELS ────────────────────────────────────────────────────────────────────
class Client(BaseModel):
    company_name: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    industry: Optional[str] = None
    address: Optional[str] = None
    status: Optional[str] = "Active"
    kra_pin: Literal[True, False] = False

class User(BaseModel):
    full_name: str
    email: str
    password: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

class UserUpdate(BaseModel):
    full_name: str
    email: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

class LoginRequest(BaseModel):
    email: str
    password: str

class PasswordResetRequest(BaseModel):
    email: str

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

class ColumnMapping(BaseModel):
    client_id: str
    file_type: Optional[str] = "general"
    original_column: str
    mapped_to: str
    confirmed_by: Optional[str] = None

class Engagement(BaseModel):
    client_id: int
    engagement_name: str
    financial_year: str
    status: Optional[str] = "Planning"
    start_date: Optional[str] = None
    end_date: Optional[str] = None

class EngagementTeam(BaseModel):
    engagement_id: int
    user_id: int
    role: str

class AuditSection(BaseModel):
    engagement_id: int
    section_name: str
    status: Optional[str] = "Pending"
    assigned_to: Optional[int] = None

class Submission(BaseModel):
    engagement_id: int
    section_id: int
    submitted_by: int
    status: Optional[str] = "Draft"
    current_stage: Optional[str] = "Accountant"
    notes: Optional[str] = None

class SubmissionStatus(BaseModel):
    status: Literal["Draft", "Submitted", "Under Review", "Changes Requested", "Approved", "Cancelled"]
    current_stage: Optional[str] = None
    notes: Optional[str] = None
    updated_by: Optional[int] = None

class Notification(BaseModel):
    user_id: int
    message: str
    type: Optional[str] = "engagement_alert"

@app.get("/")
def root():
    return {"message": "Audit AI API is running"}

@app.post("/upload")
async def upload_file_ai(
    file: UploadFile = File(...),
    client_id: str = Form(...)
):
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only.")
    MAX_FILE_SIZE = 50
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File size exceeds the maximum limit of {MAX_FILE_SIZE} MB. Uploaded file size: {file_size_mb:.2f} MB.")
    file.file.seek(0)
    file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    if ext == "pdf":
        df, source = extract_pdf(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"PDF uploaded — extracted via {source}"}

    if ext == "docx":
        df, source = extract_docx(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"DOCX uploaded — extracted via {source}"}

    try:
        df = pd.read_csv(save_path, dtype=str) if ext == "csv" else pd.read_excel(save_path, dtype=str)
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
        # Drop any reserved internal column (e.g. _row_id from a previously
        # exported workbook re-uploaded by mistake through the normal upload
        # flow) before it's ever surfaced for column detection.
        df = df.drop(columns=[c for c in df.columns if c in RESERVED_INTERNAL_COLUMNS], errors='ignore')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    save_upload(file_id, client_id, file.filename, ext, len(df))
    fill_rates = calculate_fill_rates(df)
    return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": "table",
            "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
            "fingerprint": compute_schema_fingerprint(list(df.columns)),
            "preview": df.head(5).fillna("").to_dict(orient="records"), "message": "File uploaded and processed successfully"}

@app.post("/detect-columns")
async def detect_columns_endpoint(
    client_id: str = Form(...),
    file_id: str = Form(...),
    columns: str = Form(...),
    file_type: str = Form("general"),
    fill_rates: str = Form("{}"),
    fingerprint: str = Form("")
):
    try:
        columns_list = json.loads(columns)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid columns format.")

    try:
        fill_rates_dict = json.loads(fill_rates)
    except json.JSONDecodeError:
        fill_rates_dict = {}

    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")

    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        sample_values = {}
        for col in columns_list:
            if col in df.columns:
                non_empty = df[col].dropna().replace("", float("nan")).dropna()
                sample_values[col] = str(non_empty.iloc[0]) if len(non_empty) > 0 else ""
            else:
                sample_values[col] = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    try:
        file_type_suggestion = suggest_file_type(columns_list, sample_values)
    except Exception:
        file_type_suggestion = {"file_type": "other", "file_type_label": "Other"}

    effective_file_type = file_type_suggestion["file_type"]

    computed_fingerprint = compute_schema_fingerprint(columns_list)
    save_fingerprint(client_id, computed_fingerprint, effective_file_type, columns_list)

    saved_mapping = get_mapping(client_id, effective_file_type)
    existing_fp = get_fingerprint(client_id, computed_fingerprint, effective_file_type)

    if existing_fp and saved_mapping:
        filtered_mapping = {
            col: saved_mapping.get(col)
            for col in columns_list
            if col in saved_mapping
        }
        # Require that EVERY uploaded column actually matched something in the saved
        # mapping before trusting this as a real cache hit. A fingerprint match alone
        # isn't enough proof — it's possible for this exact fingerprint to have been
        # seen before (e.g. from an earlier accidental upload of an already-cleaned/
        # exported file) while the saved MAPPING is still keyed by a completely
        # different set of column names (the original raw headers). In that case
        # filtered_mapping would be empty or only partially filled, and returning it
        # as-is silently produces a mapping table with missing or zero rows instead
        # of a real detection result. Falling through to full LLM detection below is
        # the safe behavior whenever the match isn't complete.
        if filtered_mapping and len(filtered_mapping) == len(columns_list):
            result = build_detection_result(columns_list, filtered_mapping, sample_values, fill_rates_dict)
            result["file_id"] = file_id
            result["source"] = "fingerprint_cache"
            result["message"] = "Mapping reused from cache."
            result["suggested_file_type"] = file_type_suggestion["file_type"]
            result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
            return result

    if saved_mapping:
        all_mapped = all(col in saved_mapping for col in columns_list)
        if all_mapped:
            filtered_mapping = {col: saved_mapping[col] for col in columns_list}
            result = build_detection_result(columns_list, filtered_mapping, sample_values, fill_rates_dict)
            result["file_id"] = file_id
            result["source"] = "saved_mapping"
            result["message"] = "Mapping loaded from saved client profile — LLM skipped."
            result["suggested_file_type"] = file_type_suggestion["file_type"]
            result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
            return result

    try:
        mapping = detect_columns_with_llm(columns_list, sample_values, fill_rates_dict)
        if not mapping:
            raise HTTPException(status_code=500, detail="LLM returned empty mapping.")

        result = build_detection_result(columns_list, mapping, sample_values, fill_rates_dict)

        dedup_warnings = []
        seen_targets = set()
        for original_col, info in list(result["mapping"].items()):
            if not isinstance(info, dict):
                continue
            target = str(info.get("mapped_to", "")).strip()
            if target in ("", "unknown"):
                continue
            if target in seen_targets:
                suggestion = target
                result["mapping"][original_col] = {
                    "mapped_to": "unknown",
                    "field_type": "unknown",
                    "suggestion": suggestion,
                    "sample_value": info.get("sample_value", ""),
                    "fill_rate": info.get("fill_rate", 1.0),
                }
                dedup_warnings.append(
                    f"'{original_col}' was also mapped to '{target}' — demoted to unknown. "
                    f"Please give it a unique 'Mapped To' name."
                )
            else:
                seen_targets.add(target)

        unknown_columns = [
            col for col, info in result["mapping"].items()
            if info.get("mapped_to") == "unknown" or info.get("field_type") == "unknown"
        ]
        result["unknown_count"] = len(unknown_columns)
        result["requires_manual_mapping"] = len(unknown_columns) > 0
        if dedup_warnings:
            if "warnings" in result and isinstance(result["warnings"], list):
                result["warnings"].extend(
                    {"type": "dedup_mapping", "message": w, "action": "Please update the mapped-to field to a unique name."}
                    for w in dedup_warnings
                )
            else:
                result["warnings"] = [
                    {"type": "dedup_mapping", "message": w, "action": "Please update the mapped-to field to a unique name."}
                    for w in dedup_warnings
                ]
        if result.get("warnings"):
            result["warning_count"] = len(result["warnings"])

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")

    result["file_id"] = file_id
    result["source"] = "llm_detection"
    result["suggested_file_type"] = file_type_suggestion["file_type"]
    result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
    return result

@app.post("/save-mapping")
async def save_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    mapping: str = Form(...),
    confirmed_by: str = Form(None)
):
    try:
        mapping_dict = json.loads(mapping)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid mapping format.")
    if not mapping_dict:
        raise HTTPException(status_code=400, detail="Mapping cannot be empty. Please provide a valid mapping.")
    save_mapping(client_id, file_type, mapping_dict, confirmed_by)
    return {"client_id": client_id, "file_type": file_type, "columns_saved": len(mapping_dict),
            "message": f"Mapping saved successfully for client {client_id} and file type {file_type}."}

@app.get("/get-mapping/{client_id}")
async def get_mapping_endpoint(client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        return {"client_id": client_id, "file_type": file_type, "mapping": {},
                "message": "No saved mapping found for this client."}
    return {"client_id": client_id, "file_type": file_type, "mapping": mapping,
            "columns_mapped": len(mapping), "message": "Saved mapping retrieved successfully."}

@app.get("/uploads/{client_id}")
async def get_uploads_endpoint(client_id: str):
    uploads = get_uploads(client_id)
    return {"client_id": client_id, "total_uploads": len(uploads), "uploads": uploads}

@app.post("/clean")
async def clean_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
) -> dict:
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client. Please detect the columns and confirm the mapping first.")
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {"file_id": file_id, "client_id": client_id, "file_type": file_type,
            "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
            "validation_report": report,
            "can_proceed": report.get("can_proceed", False),
            "message": "File cleaned successfully."}

@app.get("/clean/export-cleaned/{file_id}")
async def export_cleaned_workbook(file_id: str, client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client. Please detect the columns and confirm the mapping first.")
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)
    workbook_buffer = build_cleaning_workbook(cleaned_df, report, mapping)
    download_filename = f"{file_id}_cleaned_data.xlsx"
    return StreamingResponse(
        workbook_buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{download_filename}"'}
    )

@app.post("/clean/submit-corrected-excel")
async def submit_corrected_excel(
    file: UploadFile = File(...),
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    corrected_by: str = Form(None),
):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client. Please detect the columns and confirm the mapping first.")

    snapshot_rows = get_cleaning_snapshot(file_id, client_id, file_type)
    if not snapshot_rows:
        raise HTTPException(
            status_code=400,
            detail="No downloaded snapshot found for this file. Please download the cleaned Excel first, edit it, then upload it back."
        )

    temp_path = os.path.join(UPLOAD_DIR, f"corrected_{uuid.uuid4()}.xlsx")
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        diff_result = diff_uploaded_against_snapshot(temp_path, snapshot_rows)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        # Clean up the temp file, but never let a cleanup failure crash the request.
        # On Windows, openpyxl can briefly hold a file handle even after close() in some
        # edge cases (antivirus scanning, OS-level delays) — a delete failure here should
        # not lose the auditor's corrections, which by this point may already be processed.
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except (PermissionError, OSError):
            pass

    corrections = diff_result["corrections"]
    deleted_row_ids = diff_result["deleted_row_ids"]
    deleted_columns = diff_result["deleted_columns"]
    renamed_columns = diff_result["renamed_columns"]
    ambiguous_changes = diff_result["ambiguous_changes"]

    # More than one column's identity changed at once (more than one name missing
    # and/or more than one name added). We don't guess which old name maps to which
    # new name in this case — doing so by position previously misattributed renames
    # to completely unrelated columns whenever the schema had drifted between the
    # snapshot and this upload (e.g. resolving a column elsewhere introduced a new
    # derived field). Ask for a fresh workbook instead of risking silent data loss.
    if ambiguous_changes:
        raise HTTPException(
            status_code=400,
            detail=(
                "Multiple column changes were detected at once, so the exact rename(s) "
                "couldn't be determined safely. This usually happens when the downloaded "
                "workbook is out of date relative to the current data (for example, if "
                "another correction round already changed the columns since this file was "
                "downloaded). Please download a fresh copy of the workbook, reapply your "
                "edits there, and upload that instead."
            )
        )

    # A single, unambiguous rename was detected. Before applying it, make sure the new
    # name doesn't collide with a `mapped_to` value already used by a DIFFERENT column —
    # that would create two columns sharing one name after rename_columns runs, which
    # crashes deep inside cleaning (the exact problem detect_duplicate_mappings guards
    # against for the initial mapping). Reject clearly here instead of letting that happen.
    if renamed_columns:
        old_name, new_name = next(iter(renamed_columns.items()))
        colliding_column = next(
            (
                original_col for original_col, info in mapping.items()
                if original_col != old_name
                and isinstance(info, dict)
                and info.get("mapped_to") == new_name
            ),
            None,
        )
        if colliding_column:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Cannot rename '{old_name}' to '{new_name}': that name is already used "
                    f"by column '{colliding_column}'. Please choose a different name for "
                    f"one of these two columns and re-upload."
                )
            )

    if not corrections and not deleted_row_ids and not deleted_columns and not renamed_columns:
        raise HTTPException(
            status_code=400,
            detail="No changes were detected in the uploaded file compared to what was downloaded."
        )

    if corrections:
        save_cleaning_corrections(file_id, client_id, file_type, corrections, corrected_by)

    if deleted_row_ids:
        deletion_records = [
            {"row_index": row_id, "column": "_row_deleted",
             "original_value": "present", "corrected_value": "deleted_by_auditor"}
            for row_id in deleted_row_ids
        ]
        save_cleaning_corrections(file_id, client_id, file_type, deletion_records, corrected_by)

    if deleted_columns:
        column_deletion_records = [
            {"row_index": -1, "column": f"_column_deleted:{col}",
             "original_value": "present", "corrected_value": "deleted_by_auditor"}
            for col in deleted_columns
        ]
        save_cleaning_corrections(file_id, client_id, file_type, column_deletion_records, corrected_by)

    # Apply the rename directly to the saved mapping — same treatment whether the
    # column was previously "unknown" or already fully resolved. The auditor typing a
    # new header in Excel is the same deliberate action either way, and we've already
    # confirmed above it's unambiguous and collision-free, so there's no need to gate
    # it further or route it through a separate mapping screen.
    #
    # IMPORTANT: `old_name` from the diff is the CLEANED-DATA column name the auditor
    # saw in the workbook — for an unresolved column that's still its original raw
    # name (e.g. "DEPARTMENT"), but for an ALREADY-RESOLVED column it's the `mapped_to`
    # value (e.g. "department"), NOT the mapping dict's key (e.g. "dept"). Naively doing
    # `mapping[old_name] = ...` for a resolved column creates a bogus new entry under a
    # key that never existed in the source file, while leaving the real entry untouched
    # — silently failing to apply the rename at all. We must look up which original
    # column currently has mapped_to == old_name first.
    if renamed_columns:
        old_name, new_name = next(iter(renamed_columns.items()))

        original_key = next(
            (oc for oc, info in mapping.items()
             if isinstance(info, dict) and info.get("mapped_to") == old_name),
            None,
        )
        if original_key is None:
            # Not found as a mapped_to value anywhere — this means old_name IS itself
            # the original raw column key (the unresolved-column case, where mapped_to
            # was "unknown" and the column kept its original name throughout cleaning).
            original_key = old_name

        updated_mapping = dict(mapping)
        existing_info = updated_mapping.get(original_key, {})
        updated_mapping[original_key] = {
            **(existing_info if isinstance(existing_info, dict) else {}),
            "mapped_to": new_name,
            # Keep the existing field_type if one was already set (e.g. this column was
            # already resolved as text/date/numeric) — a rename alone doesn't change what
            # kind of data is in the column. Only default to "text" if it was genuinely
            # unset/unknown, since the auditor didn't get to pick a type through this path.
            "field_type": (
                existing_info.get("field_type")
                if isinstance(existing_info, dict) and existing_info.get("field_type") not in (None, "unknown")
                else "text"
            ),
        }
        save_mapping(client_id, file_type, updated_mapping, corrected_by)
        mapping = updated_mapping

    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    # Refresh the snapshot so the NEXT diff round compares against today's actual
    # state, not an increasingly stale baseline — this is what keeps future renames
    # and corrections unambiguous instead of drifting into ambiguous_changes territory.
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)

    rename_summary = (
        f", column '{next(iter(renamed_columns.keys()))}' renamed to "
        f"'{next(iter(renamed_columns.values()))}' in the mapping"
        if renamed_columns else ""
    )

    return {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "corrections_applied": len(corrections),
        "rows_deleted": deleted_row_ids,
        "columns_deleted": deleted_columns,
        "columns_renamed": renamed_columns,
        "message": (
            f"{len(corrections)} correction(s), {len(deleted_row_ids)} row deletion(s), "
            f"and {len(deleted_columns)} column deletion(s) applied from the uploaded file"
            f"{rename_summary}. File re-cleaned successfully."
        )
    }

@app.post("/clean/acknowledge-issue")
async def acknowledge_issue_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    issue: str = Form(...),
    acknowledged_by: str = Form(None),
):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client.")
    try:
        issue_dict = json.loads(issue)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid issue format.")
    issue_id = issue_dict.get("issue_id") or issue_fingerprint(file_id, client_id, file_type, issue_dict)
    save_cleaning_acknowledgment(issue_id, file_id, client_id, file_type, issue_dict, acknowledged_by)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {
        "file_id": file_id, "client_id": client_id, "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "message": "Issue acknowledged."
    }

@app.post("/clean/submit-inline-corrections")
async def submit_inline_corrections_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    corrections: str = Form(...),
    corrected_by: str = Form(None),
):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this client.")
    try:
        corrections_list = json.loads(corrections)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid corrections format.")
    if not corrections_list:
        raise HTTPException(status_code=400, detail="No corrections provided.")
    save_cleaning_corrections(file_id, client_id, file_type, corrections_list, corrected_by)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {
        "file_id": file_id, "client_id": client_id, "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "message": f"{len(corrections_list)} correction(s) saved and re-cleaned."
    }

# ── CLIENTS ───────────────────────────────────────────────────────────────────
@app.get("/clients")
def get_clients(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients")
    return cursor.fetchall()

@app.get("/clients/{client_id}")
def get_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients WHERE client_id = %s", (client_id,))
    client = cursor.fetchone()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client

@app.post("/clients")
def create_client(c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO clients (company_name, contact_person, email, phone, industry, address, status, kra_pin)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin)
    )
    db.commit()
    return {"client_id": cursor.lastrowid, "message": "Client created"}

@app.put("/clients/{client_id}")
def update_client(client_id: int, c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE clients SET company_name=%s, contact_person=%s, email=%s,
           phone=%s, industry=%s, address=%s, status=%s, kra_pin=%s WHERE client_id=%s""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin, client_id)
    )
    db.commit()
    return {"message": "Client updated"}

@app.delete("/clients/{client_id}")
def delete_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    try:
        cursor.execute("SELECT engagement_id FROM engagements WHERE client_id = %s", (client_id,))
        rows = cursor.fetchall()
        engagement_ids = [r['engagement_id'] for r in rows] if rows else []
        if engagement_ids:
            placeholders = ','.join(['%s'] * len(engagement_ids))
            cursor.execute(f"DELETE FROM submissions WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM audit_sections WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagement_team WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagements WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
        cursor.execute("DELETE FROM uploads WHERE client_id = %s", (str(client_id),))
        cursor.execute("UPDATE users SET assigned_client_id = NULL WHERE assigned_client_id = %s", (client_id,))
        cursor.execute("DELETE FROM clients WHERE client_id = %s", (client_id,))
        if cursor.rowcount == 0:
            db.rollback()
            raise HTTPException(status_code=404, detail="Client not found")
        db.commit()
        return {"message": "Client and dependent records deleted"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Could not delete client: {str(e)}")

# ── USERS ─────────────────────────────────────────────────────────────────────
@app.get("/users")
def get_users(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT u.user_id, u.full_name, u.email, u.phone, u.role, u.status,
               u.assigned_client_id, u.created_at, c.company_name
        FROM users u LEFT JOIN clients c ON u.assigned_client_id = c.client_id
    """)
    return cursor.fetchall()

@app.get("/users/{user_id}")
def get_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE user_id = %s", (user_id,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.pop("password_hash", None)
    return user

@app.post("/users")
def create_user(u: User, db=Depends(get_db)):
    hashed = hash_password(u.password)
    cursor = db.cursor()
    try:
        cursor.execute(
            """INSERT INTO users (full_name, email, password_hash, phone, role, assigned_client_id, status)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (u.full_name, u.email, hashed, u.phone, u.role, u.assigned_client_id, u.status)
        )
        db.commit()
        return {"user_id": cursor.lastrowid, "message": "User created"}
    except Exception:
        raise HTTPException(status_code=400, detail="Email already exists")

@app.put("/users/{user_id}")
def update_user(user_id: int, u: UserUpdate, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE users SET full_name=%s, email=%s, phone=%s,
           role=%s, assigned_client_id=%s, status=%s WHERE user_id=%s""",
        (u.full_name, u.email, u.phone, u.role, u.assigned_client_id, u.status, user_id)
    )
    db.commit()
    return {"message": "User updated"}

@app.delete("/users/{user_id}")
def delete_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "User deleted"}

@app.put("/users/{user_id}/assign/{client_id}")
def assign_user_to_client(user_id: int, client_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE users SET assigned_client_id=%s WHERE user_id=%s", (client_id, user_id))
    db.commit()
    return {"message": "User assigned to client"}

# ── AUTH ──────────────────────────────────────────────────────────────────────
@app.post("/auth/login")
def login(req: LoginRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user or not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    if user["status"] != "Active":
        raise HTTPException(status_code=403, detail="Account is inactive")
    token = create_token({"user_id": user["user_id"], "email": user["email"], "role": user["role"]})
    return {"access_token": token, "token_type": "bearer",
            "user": {"user_id": user["user_id"], "full_name": user["full_name"],
                     "email": user["email"], "role": user["role"]}}

@app.post("/auth/password-reset-request")
def password_reset_request(req: PasswordResetRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    token = secrets.token_urlsafe(32)
    expires_at = datetime.utcnow() + timedelta(hours=1)
    cursor2 = db.cursor()
    cursor2.execute("INSERT INTO password_resets (user_id, token, expires_at) VALUES (%s, %s, %s)",
                    (user["user_id"], token, expires_at))
    db.commit()
    return {"message": "Password reset token generated", "token": token}

@app.post("/auth/password-reset-confirm")
def password_reset_confirm(req: PasswordResetConfirm, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM password_resets WHERE token = %s AND expires_at > NOW()", (req.token,))
    reset = cursor.fetchone()
    if not reset:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    hashed = hash_password(req.new_password)
    cursor2 = db.cursor()
    cursor2.execute("UPDATE users SET password_hash = %s WHERE user_id = %s", (hashed, reset["user_id"]))
    cursor2.execute("DELETE FROM password_resets WHERE token = %s", (req.token,))
    db.commit()
    return {"message": "Password reset successful"}

# ── COLUMN MAPPINGS ───────────────────────────────────────────────────────────
@app.get("/column-mappings")
def get_all_mappings(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings")
    return cursor.fetchall()

@app.get("/column-mappings/{client_id}")
def get_client_mappings(client_id: str, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

@app.post("/column-mappings")
def create_mapping(m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO column_mappings (client_id, file_type, original_column, mapped_to, confirmed_by)
           VALUES (%s, %s, %s, %s, %s)""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by)
    )
    db.commit()
    return {"id": cursor.lastrowid, "message": "Column mapping created"}

@app.put("/column-mappings/{mapping_id}")
def update_mapping(mapping_id: int, m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE column_mappings SET client_id=%s, file_type=%s, original_column=%s,
           mapped_to=%s, confirmed_by=%s WHERE id=%s""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by, mapping_id)
    )
    db.commit()
    return {"message": "Column mapping updated"}

@app.delete("/column-mappings/{mapping_id}")
def delete_mapping(mapping_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM column_mappings WHERE id = %s", (mapping_id,))
    db.commit()
    return {"message": "Column mapping deleted"}

# ── ENGAGEMENTS ───────────────────────────────────────────────────────────────
@app.get("/engagements")
def get_engagements(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        ORDER BY e.created_at DESC
    """)
    return cursor.fetchall()

@app.get("/engagements/{engagement_id}")
def get_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    return engagement

@app.post("/engagements")
def create_engagement(e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO engagements (client_id, engagement_name, financial_year, status, start_date, end_date)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date)
    )
    engagement_id = cursor.lastrowid
    for section in ["Revenue", "Expenses", "Inventory", "Cash & Bank"]:
        cursor.execute("INSERT INTO audit_sections (engagement_id, section_name) VALUES (%s, %s)",
                       (engagement_id, section))
    db.commit()
    return {"engagement_id": engagement_id, "message": "Engagement created with default audit sections"}

@app.put("/engagements/{engagement_id}")
def update_engagement(engagement_id: int, e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE engagements SET client_id=%s, engagement_name=%s, financial_year=%s,
           status=%s, start_date=%s, end_date=%s WHERE engagement_id=%s""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date, engagement_id)
    )
    db.commit()
    return {"message": "Engagement updated"}

@app.delete("/engagements/{engagement_id}")
def delete_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagements WHERE engagement_id = %s", (engagement_id,))
    db.commit()
    return {"message": "Engagement deleted"}

# ── ENGAGEMENT TEAM ───────────────────────────────────────────────────────────
@app.get("/engagements/{engagement_id}/team")
def get_engagement_team(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT et.*, u.full_name, u.email, u.role FROM engagement_team et
        LEFT JOIN users u ON et.user_id = u.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

@app.post("/engagements/{engagement_id}/team")
def add_team_member(engagement_id: int, t: EngagementTeam, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("INSERT INTO engagement_team (engagement_id, user_id, role) VALUES (%s, %s, %s)",
                   (engagement_id, t.user_id, t.role))
    db.commit()
    return {"team_id": cursor.lastrowid, "message": "Team member added"}

@app.delete("/engagements/{engagement_id}/team/{user_id}")
def remove_team_member(engagement_id: int, user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id=%s AND user_id=%s",
                   (engagement_id, user_id))
    db.commit()
    return {"message": "Team member removed"}

# ── AUDIT SECTIONS ────────────────────────────────────────────────────────────
@app.get("/engagements/{engagement_id}/sections")
def get_audit_sections(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as assigned_to_name FROM audit_sections s
        LEFT JOIN users u ON s.assigned_to = u.user_id
        WHERE s.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

@app.post("/engagements/{engagement_id}/sections")
def add_audit_section(engagement_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "INSERT INTO audit_sections (engagement_id, section_name, status, assigned_to) VALUES (%s, %s, %s, %s)",
        (engagement_id, s.section_name, s.status, s.assigned_to)
    )
    db.commit()
    return {"section_id": cursor.lastrowid, "message": "Audit section added"}

@app.put("/audit-sections/{section_id}")
def update_audit_section(section_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "UPDATE audit_sections SET section_name=%s, status=%s, assigned_to=%s WHERE section_id=%s",
        (s.section_name, s.status, s.assigned_to, section_id)
    )
    db.commit()
    return {"message": "Audit section updated"}

@app.delete("/audit-sections/{section_id}")
def delete_audit_section(section_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE section_id = %s", (section_id,))
    db.commit()
    return {"message": "Audit section deleted"}

# ── SUBMISSIONS ───────────────────────────────────────────────────────────────
@app.get("/audit-sections/{section_id}/latest-submission")
def get_section_latest_submission(section_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        "SELECT s.*, u.full_name as submitted_by_name "
        "FROM submissions s "
        "LEFT JOIN users u ON s.submitted_by = u.user_id "
        "WHERE s.section_id = %s "
        "ORDER BY s.created_at DESC LIMIT 1",
        (section_id,)
    )
    row = cursor.fetchone()
    return row if row else None

@app.get("/submissions")
def get_all_submissions(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        ORDER BY s.created_at DESC
    """)
    return cursor.fetchall()

@app.get("/submissions/{submission_id}")
def get_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.submission_id = %s
    """, (submission_id,))
    submission = cursor.fetchone()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    return submission

@app.post("/submissions")
def create_submission(s: Submission, db=Depends(get_db)):
    insert_cursor = db.cursor()
    insert_cursor.execute(
        """INSERT INTO submissions (engagement_id, section_id, submitted_by, status, current_stage, notes)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (s.engagement_id, s.section_id, s.submitted_by, s.status, s.current_stage, s.notes)
    )
    submission_id = insert_cursor.lastrowid
    cursor = db.cursor(dictionary=True)
    if s.current_stage and s.current_stage != "Accountant":
        cursor.execute("""
            SELECT e.engagement_name, sec.section_name FROM engagements e
            LEFT JOIN audit_sections sec ON sec.engagement_id = e.engagement_id
            WHERE sec.section_id = %s
        """, (s.section_id,))
        info = cursor.fetchone()
        if info:
            message = f"{info['section_name']} for {info['engagement_name']} is now {s.status}"
            cursor.execute("""
                SELECT u.user_id FROM users u
                INNER JOIN engagement_team et ON u.user_id = et.user_id
                WHERE et.engagement_id = %s AND u.role = %s
            """, (s.engagement_id, s.current_stage))
            for auditor in cursor.fetchall():
                cursor.execute("INSERT INTO notifications (user_id, message, type) VALUES (%s, %s, %s)",
                               (auditor['user_id'], message, 'engagement_alert'))
    db.commit()
    return {"submission_id": submission_id, "message": "Submission created"}

@app.put("/submissions/{submission_id}/status")
def update_submission_status(submission_id: int, s: SubmissionStatus, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT sub.*, e.engagement_name, e.engagement_id, sec.section_name
        FROM submissions sub
        LEFT JOIN engagements e ON sub.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON sub.section_id = sec.section_id
        WHERE sub.submission_id = %s
    """, (submission_id,))
    sub = cursor.fetchone()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")

    cursor2 = db.cursor(dictionary=True)
    if s.updated_by:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s, submitted_by=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, s.updated_by, submission_id)
        )
    else:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, submission_id)
        )

    target_roles = []
    if s.current_stage:
        target_roles = [s.current_stage]
    elif s.status in ("Approved", "Cancelled"):
        target_roles = ["Accountant", "Auditor", "Senior Auditor", "Assistant Manager",
                        "Audit Manager", "Engagement Partner", "Quality Reviewer"]

    if target_roles:
        message = f"{sub['section_name']} for {sub['engagement_name']} is now {s.status}"
        cursor2.execute(f"""
            SELECT u.user_id FROM users u
            INNER JOIN engagement_team et ON u.user_id = et.user_id
            WHERE et.engagement_id = %s AND u.role IN ({','.join(['%s']*len(target_roles))})
        """, (sub['engagement_id'], *target_roles))
        for row in cursor2.fetchall():
            cursor2.execute("INSERT INTO notifications (user_id, message, type) VALUES (%s, %s, %s)",
                            (row['user_id'], message, "engagement_alert"))

    db.commit()
    return {"message": f"Submission status updated to {s.status}"}

@app.delete("/submissions/{submission_id}")
def delete_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM submissions WHERE submission_id = %s", (submission_id,))
    db.commit()
    return {"message": "Submission deleted"}

# ── NOTIFICATIONS ─────────────────────────────────────────────────────────────
@app.get("/notifications/{user_id}")
def get_user_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
    return cursor.fetchall()

@app.get("/notifications/{user_id}/unread")
def get_unread_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s AND is_read = FALSE ORDER BY created_at DESC",
                   (user_id,))
    return cursor.fetchall()

@app.put("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE notification_id = %s", (notification_id,))
    db.commit()
    return {"message": "Notification marked as read"}

@app.put("/notifications/{user_id}/read-all")
def mark_all_read(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "All notifications marked as read"}

# ── FILE UPLOAD ───────────────────────────────────────────────────────────────
@app.post("/clients/{client_id}/upload")
def upload_client_file(client_id: int, file: UploadFile = File(...), db=Depends(get_db)):
    allowed_types = ["xlsx", "xls", "csv", "pdf", "tiff", "tif", "jpg", "jpeg", "png", "xml", "json", "txt"]
    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail="File format not allowed.")
    file_path = f"{UPLOAD_DIR}/{client_id}_{file.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    cursor = db.cursor()
    cursor.execute("INSERT INTO uploads (client_id, file_name, file_type, file_path) VALUES (%s, %s, %s, %s)",
                   (client_id, file.filename, file_ext.upper(), file_path))
    db.commit()
    return {"file_id": cursor.lastrowid, "filename": file.filename, "type": file_ext.upper(),
            "message": "File uploaded successfully"}

@app.get("/clients/{client_id}/files")
def get_client_files(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM uploads WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

@app.get("/files")
def get_all_files(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT f.*, c.company_name FROM uploads f
        LEFT JOIN clients c ON f.client_id = c.client_id
        ORDER BY f.upload_date DESC
    """)
    return cursor.fetchall()

# ── RUN ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("Audit:app", host="0.0.0.0", port=8000, reload=True)
